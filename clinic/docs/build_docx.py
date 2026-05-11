from docx import Document
from pathlib import Path

md_path = Path(r"C:\xampp\htdocs\clinic\docs\System_Walkthrough_AR.md")
docx_path = Path(r"C:\xampp\htdocs\clinic\docs\System_Walkthrough_AR.docx")

document = Document()

for raw_line in md_path.read_text(encoding="utf-8").splitlines():
    line = raw_line.rstrip()
    if not line:
        document.add_paragraph("")
        continue

    if line.startswith("### "):
        document.add_heading(line[4:].strip(), level=3)
        continue
    if line.startswith("## "):
        document.add_heading(line[3:].strip(), level=2)
        continue
    if line.startswith("# "):
        document.add_heading(line[2:].strip(), level=1)
        continue

    if line.startswith("- "):
        document.add_paragraph(line[2:].strip(), style="List Bullet")
        continue

    if line[:2].isdigit() and line[1:3] == ") ":
        document.add_paragraph(line[3:].strip(), style="List Number")
        continue

    document.add_paragraph(line)

document.save(docx_path)
print(str(docx_path))
